"""文件读取服务：按扩展名分派解析器，支持 TXT / MD / PDF / DOCX"""
import os

from fastapi import UploadFile

from app.core.exceptions import BusinessException

# 支持的扩展名集合
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


async def read_file(file: UploadFile) -> str:
    """
    根据上传文件名扩展名分派解析器，返回纯文本。
    不支持的格式抛出 BusinessException(400)。
    """
    filename = file.filename or ""
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise BusinessException(
            code=400,
            msg=f"暂不支持「{ext or '未知'}」格式，请上传 txt / md / pdf / docx 文件"
        )

    content_bytes = await file.read()
    await file.seek(0)

    if ext in (".txt", ".md"):
        return _decode_text(content_bytes)
    if ext == ".pdf":
        return extract_pdf(content_bytes)
    if ext == ".docx":
        return extract_docx(content_bytes)
    return _decode_text(content_bytes)


def _decode_text(content_bytes: bytes) -> str:
    """TXT / MD 解码，优先 UTF-8，失败回退 GBK"""
    for encoding in ("utf-8", "gbk", "utf-8-sig"):
        try:
            return content_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content_bytes.decode("utf-8", errors="ignore")


def extract_pdf(content_bytes: bytes) -> str:
    """PDF 提取文本（PyMuPDF），逐页拼接"""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content_bytes, filetype="pdf")
    try:
        pages = [page.get_text() for page in doc]
    finally:
        doc.close()
    text = "\n\n".join(p for p in pages if p.strip())
    if not text.strip():
        raise BusinessException(code=400, msg="PDF 未提取到文本，可能是扫描件，暂不支持 OCR")
    return text


def extract_docx(content_bytes: bytes) -> str:
    """Word (.docx) 提取段落与表格文本"""
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(content_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格内容逐行拼接
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        raise BusinessException(code=400, msg="Word 文档未提取到文本内容")
    return text


def read_local_text(file_path: str) -> str:
    """读取磁盘本地文本文件（调试用）"""
    ext = os.path.splitext(file_path)[1].lower()
    with open(file_path, "rb") as f:
        content = f.read()
    if ext == ".pdf":
        return extract_pdf(content)
    if ext == ".docx":
        return extract_docx(content)
    return _decode_text(content)
